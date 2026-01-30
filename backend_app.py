"""
SPSS Syntax Generator - UPGRADED VERSION 2.0
Automatické generování SPSS syntax z dat a dotazníku
S pokročilou podporou filtrovaných otázek
"""

import pyreadstat
import re
from pathlib import Path
from typing import Dict, List, Tuple, Any
import docx

# Import naší nové parsovací logiky
def parse_questionnaire_from_docx(docx_path: str) -> Dict:
    """
    Parsuje Word dokument s dotazníkem a extrahuje otázky.
    Vrací strukturovaná data o otázkách.
    """
    doc = docx.Document(docx_path)
    
    questions = []
    current_q = None
    collecting_items = False
    
    i = 0
    while i < len(doc.paragraphs):
        text = doc.paragraphs[i].text.strip()
        
        if not text:
            i += 1
            continue
        
        # Detekce nové otázky - kód s tečkou
        # Buď má otázku (? nebo :), nebo následuje "Vyberte typ otázky::"
        if re.match(r'^[A-Z0-9][A-Za-z0-9_]*\.', text):
            has_question_mark = ('?' in text or ':' in text)
            
            # Pokud už má otazník/dvojtečku, nemusíme hledat dopředu
            if has_question_mark:
                next_has_type = False
            else:
                # Pouze pokud NEMÁ otazník, podíváme se dopředu (optimalizace)
                next_has_type = False
                for j in range(i+1, min(i+20, len(doc.paragraphs))):
                    next_text = doc.paragraphs[j].text.strip()
                    if 'Vyberte typ otázky::' in next_text:
                        next_has_type = True
                        break
                    # Pokud narazíme na další otázku, přestaneme hledat
                    if next_text and re.match(r'^[A-Z0-9][A-Za-z0-9_]*\.', next_text):
                        break
            
            if has_question_mark or next_has_type:
                # Uložit předchozí otázku
                if current_q:
                    questions.append(current_q)
                
                code = text.split('.')[0].strip()
                current_q = {
                    'code': code,
                    'text': text,
                    'type': None,
                    'items': []
                }
                collecting_items = True
        
        # Detekce typu otázky
        elif 'Vyberte typ otázky::' in text:
            if current_q:
                qtype = text.replace('Vyberte typ otázky::', '').strip()
                # Pokud ještě nemá typ, nebo je to relevantnější typ
                if current_q['type'] is None or qtype in ['BATERIE OTÁZEK - JEDNA MOŽNÁ ODPOVĚĎ', 'VÍCE MOŽNÝCH ODPOVĚDÍ', 'FILTRACE ODPOVĚDÍ', 'FILTRACE ODPOVĚDÍ BATERIE - JEDNA MOŽNÁ ODPOVĚĎ', 'FILTRACE ODPOVĚDÍ BATERIE MULTIPLE']:
                    current_q['type'] = qtype
                collecting_items = False
        
        # Přeskakujeme metadata
        elif any(keyword in text for keyword in [
            'Nastavení otázky', 'Povinná', 'Zvolených minimálně',
            'Pravidla', 'IF (', 'THEN', 'Min.', 'Max.', 
            'Délka textu', 'Minimální hodnota', 'Jdi na', 'Pokud uživatel'
        ]):
            collecting_items = False
        
        # Sbíráme položky a stupnice
        elif current_q and collecting_items and text and not text.startswith('<img') and not text.startswith('#'):
            current_q['items'].append(text)
        
        i += 1
    
    # Uložit poslední otázku
    if current_q:
        questions.append(current_q)
    
    # Kategorizace otázek - BEZ rozdělování na items/scales
    multiple_response = []
    batteries = []
    filtered_multiple = []
    filtered_batteries = []
    filtered_batteries_multiple = []
    
    for q in questions:
        if q['type'] == 'VÍCE MOŽNÝCH ODPOVĚDÍ' and len(q['items']) > 0:
            multiple_response.append(q)
        elif q['type'] == 'BATERIE OTÁZEK - JEDNA MOŽNÁ ODPOVĚĎ' and len(q['items']) > 0:
            batteries.append(q)
        elif q['type'] == 'FILTRACE ODPOVĚDÍ' and len(q['items']) > 0:
            filtered_multiple.append(q)
        elif q['type'] == 'FILTRACE ODPOVĚDÍ BATERIE - JEDNA MOŽNÁ ODPOVĚĎ' and len(q['items']) > 0:
            filtered_batteries.append(q)
        elif q['type'] == 'FILTRACE ODPOVĚDÍ BATERIE MULTIPLE' and len(q['items']) > 0:
            filtered_batteries_multiple.append(q)
    
    return {
        'multiple_response': multiple_response,
        'batteries': batteries,
        'filtered_multiple': filtered_multiple,
        'filtered_batteries': filtered_batteries,
        'filtered_batteries_multiple': filtered_batteries_multiple,
        'all_questions': questions
    }


def find_parent_question(question_code: str, all_questions: List[Dict]) -> Dict:
    """Najde rodičovskou otázku pro filtrovanou otázku."""
    for i, q in enumerate(all_questions):
        if q['code'] == question_code:
            for j in range(i-1, -1, -1):
                if all_questions[j]['type'] == 'VÍCE MOŽNÝCH ODPOVĚDÍ':
                    return all_questions[j]
    return None


class SPSSSyntaxGenerator:
    """Generátor SPSS syntax z exportovaných dat a dotazníku - UPGRADED"""
    
    def __init__(self, data_path: str, questionnaire_path: str):
        self.data_path = data_path
        self.questionnaire_path = questionnaire_path
        self.df = None
        self.meta = None
        self.questionnaire_data = None
        self.syntax_parts = []
        
    def load_data(self):
        """Načtení SPSS dat"""
        print("📂 Načítám SPSS data...")
        self.df, self.meta = pyreadstat.read_sav(self.data_path)
        print(f"   ✓ Načteno {len(self.df)} respondentů, {len(self.df.columns)} proměnných")
        
    def load_questionnaire(self):
        """Načtení dotazníku - UPGRADED s novou logikou"""
        print("📋 Načítám dotazník (UPGRADED parsing)...")
        self.questionnaire_data = parse_questionnaire_from_docx(self.questionnaire_path)
        
        total = (len(self.questionnaire_data['multiple_response']) + 
                len(self.questionnaire_data['batteries']) +
                len(self.questionnaire_data['filtered_multiple']) +
                len(self.questionnaire_data['filtered_batteries']) +
                len(self.questionnaire_data['filtered_batteries_multiple']))
        
        print(f"   ✓ Multiple Response: {len(self.questionnaire_data['multiple_response'])}")
        print(f"   ✓ Baterie: {len(self.questionnaire_data['batteries'])}")
        print(f"   ✓ Filtrované Multiple: {len(self.questionnaire_data['filtered_multiple'])}")
        print(f"   ✓ Filtrované Baterie: {len(self.questionnaire_data['filtered_batteries'])}")
        print(f"   ✓ Filt. Bat. Multiple: {len(self.questionnaire_data['filtered_batteries_multiple'])}")
        print(f"   ✓ CELKEM: {total} otázek")
        
    def get_variables_for_question(self, question_code: str) -> List[str]:
        """Získá všechny proměnné pro daný kód otázky."""
        prefix = f'Q{question_code}__'
        return [col for col in self.df.columns if col.startswith(prefix)]
    
    def get_item_text_from_label(self, label: str) -> str:
        """Extrahuje text položky z variable labelu."""
        if '|' in label:
            return label.split('|')[-1].strip()
        parts = label.split('\n')
        if len(parts) > 1:
            return parts[-1].strip()
        return label
    
    def generate_syntax(self) -> str:
        """Hlavní metoda pro generování syntax - UPGRADED"""
        print("\n🔧 Generuji SPSS syntax (UPGRADED)...")
        
        self.syntax_parts = []
        
        # 1. Filtr na dokončené respondenty
        self.syntax_parts.append("* OMEZENÍ DAT NA RESPONDENTY, KTEŘÍ DOKONČILI DOTAZNÍK.")
        self.syntax_parts.append("SELECT IF resstatus = 2.")
        self.syntax_parts.append("EXECUTE.")
        self.syntax_parts.append("FREQUENCIES resstatus.")
        self.syntax_parts.append("")
        
        # 2. Baterie otázek
        self._generate_batteries()
        
        # 3. Multiple Response  
        self._generate_multiple_response()
        
        # 4. Filtrované Multiple Response
        self._generate_filtered_multiple()
        
        # 5. Filtrované Baterie
        self._generate_filtered_batteries()
        
        # 6. Filtrované Baterie Multiple
        self._generate_filtered_batteries_multiple()
        
        # 7. MRSETS
        self._generate_mrsets()
        
        syntax = '\n'.join(self.syntax_parts)
        print(f"✅ Vygenerováno {len(syntax.split(chr(10)))} řádků syntaxu")
        return syntax
    
    def _generate_batteries(self):
        """Generuje VAR LAB pro baterie otázek"""
        if not self.questionnaire_data['batteries']:
            return
            
        section = ["* ÚPRAVA LABELŮ PRO BATERIE OTÁZEK - v tabulkách zobrazí jen text položky.", ""]
        
        for battery in self.questionnaire_data['batteries']:
            code = battery['code']
            vars_list = self.get_variables_for_question(code)
            
            if not vars_list:
                continue
            
            section.append(f"* {code} - {battery['text'][:80]}...")
            for i, item_text in enumerate(battery['items'], 1):
                var_name = f'Q{code}__{i}'
                if var_name in vars_list:
                    section.append(f'VAR LAB {var_name} "{item_text}".')
            section.append("EXECUTE.")
            section.append("")
        
        self.syntax_parts.extend(section)
    
    def _generate_multiple_response(self):
        """Generuje VAR LAB pro multiple response"""
        if not self.questionnaire_data['multiple_response']:
            return
            
        section = ["* PŘÍPRAVA MULTIPLE RESPONSE SETŮ - DICHOTOMICKÉ OTÁZKY.", ""]
        
        for mr_q in self.questionnaire_data['multiple_response']:
            code = mr_q['code']
            vars_list = self.get_variables_for_question(code)
            
            if not vars_list:
                continue
            
            # DŮLEŽITÉ: Odstranit všechny \n aby se text nezalomil bez hvězdičky
            question_text = mr_q['text'].replace('\n', ' ').strip()
            # Omezit délku
            if len(question_text) > 200:
                question_text = question_text[:197] + "..."
            
            section.append(f"* {code} - {question_text}.")
            section.append(f"* Úprava labelů na názvy jednotlivých položek.")
            
            for i, item_text in enumerate(mr_q['items'], 1):
                var_name = f'Q{code}__{i}'
                if var_name in vars_list:
                    section.append(f'VAR LAB {var_name} "{item_text}".')
            
            section.append("EXECUTE.")
            section.append("")
        
        self.syntax_parts.extend(section)
    
    def _generate_filtered_multiple(self):
        """Generuje VAR LAB pro filtrované multiple response"""
        if not self.questionnaire_data['filtered_multiple']:
            return
            
        section = ["* FILTROVANÉ MULTIPLE RESPONSE OTÁZKY.", ""]
        
        for mr_q in self.questionnaire_data['filtered_multiple']:
            code = mr_q['code']
            vars_list = self.get_variables_for_question(code)
            
            if not vars_list:
                continue
            
            # Najdeme rodičovskou otázku
            parent = find_parent_question(code, self.questionnaire_data['all_questions'])
            if not parent:
                continue
            
            section.append(f"* {code} - {mr_q['text'][:80]}...")
            section.append(f"* Používá odpovědi z {parent['code']}.")
            
            # Použijeme položky z rodiče
            for i, item_text in enumerate(parent['items'], 1):
                var_name = f'Q{code}__{parent["code"]}_{i}'
                if var_name in vars_list:
                    section.append(f'VAR LAB {var_name} "{item_text}".')
            
            # Přidáme extra odpověď
            if mr_q['items']:
                extra_var = f'Q{code}__1'
                if extra_var in vars_list:
                    section.append(f'VAR LAB {extra_var} "{mr_q["items"][0]}".')
            
            section.append("EXECUTE.")
            section.append("")
        
        self.syntax_parts.extend(section)
    
    def _generate_filtered_batteries(self):
        """Generuje VAR LAB pro filtrované baterie"""
        if not self.questionnaire_data['filtered_batteries']:
            return
            
        section = ["* FILTROVANÉ BATERIE OTÁZEK.", ""]
        
        for battery in self.questionnaire_data['filtered_batteries']:
            code = battery['code']
            vars_list = self.get_variables_for_question(code)
            
            if not vars_list:
                continue
            
            parent = find_parent_question(code, self.questionnaire_data['all_questions'])
            if not parent:
                continue
            
            section.append(f"* {code} - {battery['text'][:80]}...")
            section.append(f"* Položky jsou filtrovány z {parent['code']}.")
            
            for i, item_text in enumerate(parent['items'], 1):
                var_name = f'Q{code}__{parent["code"]}_{i}'
                if var_name in vars_list:
                    section.append(f'VAR LAB {var_name} "{item_text}".')
            
            section.append("EXECUTE.")
            section.append("")
        
        self.syntax_parts.extend(section)
    
    def _generate_filtered_batteries_multiple(self):
        """Generuje VAR LAB pro filtrované baterie multiple"""
        if not self.questionnaire_data['filtered_batteries_multiple']:
            return
            
        section = ["* FILTROVANÉ BATERIE MULTIPLE.", ""]
        
        for battery in self.questionnaire_data['filtered_batteries_multiple']:
            code = battery['code']
            all_vars = [col for col in self.df.columns if col.startswith(f'Q{code}__')]
            
            if not all_vars:
                continue
            
            parent = find_parent_question(code, self.questionnaire_data['all_questions'])
            if not parent:
                continue
            
            section.append(f"* {code} - {battery['text'][:80]}...")
            section.append(f"* Baterie multiple filtrovaná z {parent['code']}.")
            
            # Extrahujeme unique row identifiers
            rows = set()
            for var in all_vars:
                match = re.search(r'__([A-Z0-9_]+)column', var)
                if match:
                    rows.add(match.group(1))
            
            # Pro každý řádek
            for row in sorted(rows):
                match = re.search(r'_(\d+)$', row)
                if match:
                    idx = int(match.group(1))
                    if idx <= len(parent['items']):
                        item_text = parent['items'][idx - 1]
                        
                        for col_idx, col_text in enumerate(battery['items'], 1):
                            var_name = f'Q{code}__{row}column{col_idx}'
                            if var_name in all_vars:
                                full_label = f"{item_text}|{col_text}"
                                section.append(f'VAR LAB {var_name} "{full_label}".')
            
            section.append("EXECUTE.")
            section.append("")
        
        self.syntax_parts.extend(section)
    
    def _generate_mrsets(self):
        """Vytvoří MRSETS pro všechny MR otázky"""
        section = [""]
        
        # MR sety pro standardní multiple response
        for mr_q in self.questionnaire_data['multiple_response']:
            code = mr_q['code']
            vars_list = self.get_variables_for_question(code)
            
            if not vars_list:
                continue
            
            # DŮLEŽITÉ: Vyfiltrovat JEN numerické proměnné (stringové jako _jina NEPATŘÍ do MDGROUP)
            numeric_vars = [v for v in vars_list if v in self.df.columns and self.df[v].dtype in ['int64', 'float64']]
            
            if not numeric_vars:
                continue
            
            # Určíme VALUE z první numerické proměnné
            sample_var = numeric_vars[0]
            value_to_use = 2
            if sample_var in self.meta.variable_value_labels:
                val_labels = self.meta.variable_value_labels[sample_var]
                for val, label in val_labels.items():
                    if 'ano' in label.lower() or 'yes' in label.lower():
                        value_to_use = val
                        break
            
            vars_string = ' '.join(numeric_vars)
            section.append(f"* Vytvoření MR setu pro {code}.")
            section.append(f"MRSETS")
            section.append(f"  /MDGROUP NAME=${code.lower()} CATEGORYLABELS=VARLABELS ")
            section.append(f"  VARIABLES={vars_string} VALUE={value_to_use}")
            section.append(f"  /DISPLAY NAME=[${code.lower()}].")
            section.append("")
        
        # MR sety pro filtrované multiple response
        for mr_q in self.questionnaire_data['filtered_multiple']:
            code = mr_q['code']
            vars_list = self.get_variables_for_question(code)
            
            if not vars_list:
                continue
            
            # DŮLEŽITÉ: Vyfiltrovat JEN numerické proměnné
            numeric_vars = [v for v in vars_list if v in self.df.columns and self.df[v].dtype in ['int64', 'float64']]
            
            if not numeric_vars:
                continue
            
            value_to_use = 2
            if numeric_vars and numeric_vars[0] in self.meta.variable_value_labels:
                val_labels = self.meta.variable_value_labels[numeric_vars[0]]
                for val, label in val_labels.items():
                    if 'ano' in label.lower():
                        value_to_use = val
                        break
            
            vars_string = ' '.join(numeric_vars)
            section.append(f"* Vytvoření MR setu pro {code}.")
            section.append(f"MRSETS")
            section.append(f"  /MDGROUP NAME=${code.lower()} CATEGORYLABELS=VARLABELS ")
            section.append(f"  VARIABLES={vars_string} VALUE={value_to_use}")
            section.append(f"  /DISPLAY NAME=[${code.lower()}].")
            section.append("")
        
        self.syntax_parts.extend(section)
    
    def save_syntax(self, output_path: str):
        """Uloží vygenerovanou syntax do souboru"""
        syntax = self.generate_syntax()
        
        with open(output_path, 'w', encoding='cp1250', newline='\r\n') as f:
            f.write(syntax)
        print(f"\n💾 Syntax uložena do: {output_path}")
    
    def run(self, output_path: str):
        """Spustí celý proces"""
        print("="*80)
        print("SPSS SYNTAX GENERATOR 2.0 - UPGRADED")
        print("="*80)
        
        self.load_data()
        self.load_questionnaire()
        self.save_syntax(output_path)
        
        print("\n✅ HOTOVO!")
        return output_path


# Flask API
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import os
import tempfile

app = Flask(__name__)
CORS(app, resources={
    r"/api/*": {
        "origins": "*",
        "methods": ["GET", "POST", "OPTIONS"],
        "allow_headers": ["Content-Type"]
    }
})

@app.route('/api/generate', methods=['POST'])
def generate_syntax():
    """API endpoint pro generování syntax"""
    try:
        if 'sav_file' not in request.files or 'docx_file' not in request.files:
            return jsonify({'error': 'Chybí soubory'}), 400
        
        sav_file = request.files['sav_file']
        docx_file = request.files['docx_file']
        
        with tempfile.TemporaryDirectory() as tmpdir:
            sav_path = os.path.join(tmpdir, 'data.sav')
            docx_path = os.path.join(tmpdir, 'questionnaire.docx')
            output_path = os.path.join(tmpdir, 'syntax.sps')
            
            sav_file.save(sav_path)
            docx_file.save(docx_path)
            
            generator = SPSSSyntaxGenerator(sav_path, docx_path)
            generator.run(output_path)
            
            return send_file(
                output_path,
                mimetype='text/plain',
                as_attachment=True,
                download_name='generated_syntax.sps'
            )
    
    except Exception as e:
        import traceback
        error_detail = traceback.format_exc()
        print(f"ERROR: {error_detail}")
        return jsonify({'error': str(e), 'detail': error_detail}), 500

@app.route('/api/health', methods=['GET'])
def health():
    return jsonify({'status': 'ok', 'version': '2.0-upgraded'})

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=True)
